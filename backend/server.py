"""
Backend server for Rhino Transformation Builder
Integrates with kiro-cli for code transformation
"""

from flask import Flask, request, jsonify, Response
from flask_cors import CORS
import subprocess
import os
import json
import tempfile
from pathlib import Path

app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

def convert_documents_to_text(kiro_dir: Path):
    """
    Convert PDF and DOCX files in .kiro folder to text/markdown for better kiro-cli compatibility
    """
    import shutil
    
    # Convert PDFs to text
    for pdf_file in kiro_dir.glob("*.pdf"):
        txt_file = pdf_file.with_suffix('.txt')
        if not txt_file.exists():
            converted = False
            
            # Try using pdftotext if available
            if shutil.which('pdftotext'):
                try:
                    subprocess.run(['pdftotext', '-layout', str(pdf_file), str(txt_file)], 
                                 check=True, capture_output=True, timeout=30)
                    yield f"✓ Converted {pdf_file.name} to text format\n"
                    converted = True
                except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
                    pass
            
            # Fallback: Try Python pypdf if pdftotext failed
            if not converted:
                try:
                    import pypdf
                    with open(pdf_file, 'rb') as f:
                        reader = pypdf.PdfReader(f)
                        text = []
                        for page in reader.pages:
                            text.append(page.extract_text())
                        
                        with open(txt_file, 'w', encoding='utf-8') as out:
                            out.write('\n\n'.join(text))
                        
                        yield f"✓ Converted {pdf_file.name} using pypdf\n"
                        converted = True
                except Exception as e:
                    pass
            
            if not converted:
                yield f"⚠ Warning: Could not convert {pdf_file.name} - kiro-cli may not be able to read it\n"
    
    # Convert DOCX files to text
    for docx_file in kiro_dir.glob("*.docx"):
        txt_file = docx_file.with_suffix('.txt')
        if not txt_file.exists():
            converted = False
            
            # Try using python-docx library
            try:
                import docx
                doc = docx.Document(str(docx_file))
                text = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = '\t'.join(cell.text for cell in row.cells)
                        text.append(row_text)
                
                with open(txt_file, 'w', encoding='utf-8') as out:
                    out.write('\n'.join(text))
                
                yield f"✓ Converted {docx_file.name} to text format\n"
                converted = True
            except ImportError:
                yield f"⚠ Warning: python-docx not installed. Install with: pip install python-docx\n"
            except Exception as e:
                yield f"⚠ Warning: Could not convert {docx_file.name}: {str(e)}\n"
            
            if not converted:
                yield f"⚠ Warning: Could not convert {docx_file.name} - kiro-cli may not be able to read it\n"
    
    yield "Document conversion complete\n"


def run_kiro_chat(kiro_content: str, working_dir: str = None):
    """
    Executes: kiro-cli chat with the provided content in headless mode
    Streams output back to the caller
    """
    
    # Add ~/.local/bin to PATH for kiro-cli
    env = os.environ.copy()
    local_bin = os.path.expanduser("~/.local/bin")
    env["PATH"] = f"{local_bin}:{env.get('PATH', '')}"
    
    # Set working directory
    cwd = working_dir or os.getcwd()
    
    # Create a temporary spec file in the .kiro folder so file references work
    kiro_dir = Path(cwd) / ".kiro"
    kiro_dir.mkdir(parents=True, exist_ok=True)
    
    # Convert PDFs and DOCX to text format for better compatibility
    yield "Converting documents to text format...\n"
    for msg in convert_documents_to_text(kiro_dir):
        yield msg
    
    spec_file = kiro_dir / "transformation-spec.kiro"
    
    # Update file references to use .txt versions if they exist
    updated_content = kiro_content
    for pdf_file in kiro_dir.glob("*.pdf"):
        txt_file = pdf_file.with_suffix('.txt')
        if txt_file.exists():
            # Replace PDF references with TXT references
            updated_content = updated_content.replace(
                f"#[[file:.kiro/{pdf_file.name}]]",
                f"#[[file:.kiro/{txt_file.name}]]"
            )
    
    for docx_file in kiro_dir.glob("*.docx"):
        txt_file = docx_file.with_suffix('.txt')
        if txt_file.exists():
            # Replace DOCX references with TXT references
            updated_content = updated_content.replace(
                f"#[[file:.kiro/{docx_file.name}]]",
                f"#[[file:.kiro/{txt_file.name}]]"
            )
    
    # Write the spec content to file
    with open(spec_file, 'w') as f:
        f.write(updated_content)
    
    yield f"Created spec file: {spec_file}\n"
    
    # Use kiro-cli chat with file reference instead of direct content
    # This allows #[[file:...]] references to work properly
    cmd = ["kiro-cli", "chat", "--no-interactive", "--trust-all-tools", f"#[[file:.kiro/transformation-spec.kiro]]"]
    
    process = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
        env=env,
        cwd=cwd
    )
    
    # Stream output and strip ANSI escape codes
    import re
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    
    for line in iter(process.stdout.readline, ''):
        # Strip ANSI escape codes for cleaner output in browser
        clean_line = ansi_escape.sub('', line)
        yield clean_line
    
    process.wait()
    yield f"\n[Process completed with exit code: {process.returncode}]\n"


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({"status": "ok", "message": "Rhino Transformation Backend is running"})


@app.route('/api/transform', methods=['POST'])
def start_transformation():
    """
    Start the transformation process
    Expects JSON body with:
    - kiro_content: The content for kiro-cli (from module.kiro or generated spec)
    - working_dir: Optional working directory path
    """
    data = request.get_json()
    
    if not data or 'kiro_content' not in data:
        return jsonify({"error": "kiro_content is required"}), 400
    
    kiro_content = data['kiro_content']
    working_dir = data.get('working_dir', None)
    
    def generate():
        yield f"data: {json.dumps({'type': 'start', 'message': 'Starting transformation...'})}\n\n"
        yield f"data: {json.dumps({'type': 'info', 'message': 'Running kiro-cli chat --no-interactive --trust-all-tools'})}\n\n"
        yield f"data: {json.dumps({'type': 'info', 'message': '-' * 60})}\n\n"
        
        for line in run_kiro_chat(kiro_content, working_dir):
            yield f"data: {json.dumps({'type': 'output', 'message': line})}\n\n"
        
        yield f"data: {json.dumps({'type': 'complete', 'message': 'Transformation complete!'})}\n\n"
    
    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no'
        }
    )


@app.route('/api/convert-document', methods=['POST'])
def convert_document():
    """
    Convert a PDF or DOCX document to text
    Expects JSON body with:
    - fileName: Name of the file
    - fileContent: Array of bytes
    """
    data = request.get_json()
    
    if not data or 'fileName' not in data or 'fileContent' not in data:
        return jsonify({"error": "fileName and fileContent are required"}), 400
    
    file_name = data['fileName']
    file_content = bytes(data['fileContent'])
    
    try:
        text_content = None
        file_name_lower = file_name.lower()
        
        # Handle PDF conversion
        if file_name_lower.endswith('.pdf'):
            try:
                import pypdf
                import io
                
                pdf_file = io.BytesIO(file_content)
                reader = pypdf.PdfReader(pdf_file)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                
                text_content = '\n\n'.join(text)
            except Exception as e:
                return jsonify({"error": f"Failed to convert PDF: {str(e)}"}), 500
        
        # Handle DOCX conversion
        elif file_name_lower.endswith('.docx'):
            try:
                import docx
                import io
                
                docx_file = io.BytesIO(file_content)
                doc = docx.Document(docx_file)
                text = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = '\t'.join(cell.text for cell in row.cells)
                        text.append(row_text)
                
                text_content = '\n'.join(text)
            except Exception as e:
                return jsonify({"error": f"Failed to convert DOCX: {str(e)}"}), 500
        else:
            return jsonify({"error": "Unsupported file type. Only PDF and DOCX are supported."}), 400
        
        return jsonify({
            "success": True,
            "textContent": text_content,
            "fileName": file_name
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/save-spec', methods=['POST'])
def save_spec():
    """
    Save the transformation spec to a .kiro file
    Expects JSON body with:
    - content: The spec content
    - target_dir: Directory to save the file
    - filename: Optional filename (default: module.kiro)
    """
    data = request.get_json()
    
    if not data or 'content' not in data or 'target_dir' not in data:
        return jsonify({"error": "content and target_dir are required"}), 400
    
    content = data['content']
    target_dir = Path(data['target_dir'])
    filename = data.get('filename', 'module.kiro')
    
    try:
        # Create directory if it doesn't exist
        target_dir.mkdir(parents=True, exist_ok=True)
        
        # Write the file
        file_path = target_dir / filename
        with open(file_path, 'w') as f:
            f.write(content)
        
        return jsonify({
            "success": True,
            "message": f"Spec saved to {file_path}",
            "path": str(file_path)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    print("=" * 60)
    print("Rhino Transformation Backend Server")
    print("=" * 60)
    print("Endpoints:")
    print("  GET  /api/health    - Health check")
    print("  POST /api/transform - Start transformation (SSE stream)")
    print("  POST /api/save-spec - Save spec to file")
    print("=" * 60)
    app.run(host='0.0.0.0', port=5007, debug=True)

